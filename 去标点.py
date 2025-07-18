#请阅读井号后内容
#其中“###”代表使用者需要自行添加内容，共计5处

import os
import time
import pdfplumber
from tenacity import retry, stop_after_attempt, wait_exponential
from volcenginesdkarkruntime import Ark
from docx import Document
from docx.shared import Pt
#上面内容需要自行配置环境。先试运行，再将报错信息发给ai

class PDFPunctuationProcessor:
    def __init__(self, api_key, input_pdf_path, txt_output_path, doc_output_path, chunk_size=10000, delay=5):
        self.api_key = api_key
        self.input_pdf_path = input_pdf_path
        self.txt_output_path = txt_output_path
        self.doc_output_path = doc_output_path
        self.chunk_size = chunk_size
        self.delay = delay
        self.client = Ark(api_key=api_key)

    def read_pdf(self):
        paragraphs = []
        with pdfplumber.open(self.input_pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    page_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                    paragraphs.extend(page_paragraphs)
        return paragraphs

    def split_long_paragraph(self, paragraph):
        if len(paragraph) <= self.chunk_size:
            return [paragraph]
        sub_paragraphs = []
        current_start = 0
        while current_start < len(paragraph):
            end_pos = current_start + self.chunk_size
            if end_pos >= len(paragraph):
                sub_paragraphs.append(paragraph[current_start:])
                break
            period_pos = paragraph.rfind('。', current_start, end_pos)
            if period_pos == -1:
                semicolon_pos = paragraph.rfind('；', current_start, end_pos)
                if semicolon_pos != -1:
                    sub_paragraphs.append(paragraph[current_start:semicolon_pos + 1])
                    current_start = semicolon_pos + 1
                else:
                    sub_paragraphs.append(paragraph[current_start:end_pos])
                    current_start = end_pos
            else:
                sub_paragraphs.append(paragraph[current_start:period_pos + 1])
                current_start = period_pos + 1
        return sub_paragraphs

    def create_chunks(self, paragraphs):
        chunks = []
        current_chunk = ""
        for para in paragraphs:
            if len(para) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                sub_paragraphs = self.split_long_paragraph(para)
                chunks.extend(sub_paragraphs)
            else:
                if len(current_chunk) + len(para) <= self.chunk_size:
                    current_chunk += "\n\n" + para if current_chunk else para
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = para
        if current_chunk:
            chunks.append(current_chunk)
        return chunks

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def call_doubao_api(self, text):
        prompt = f"请为以下没有标点的文本添加合适的标点符号，保持原文语义不变：\n{text}"
        try:
            response = self.client.chat.completions.create(
                model="medel_id",                  ### 引号内替换为模型id(如doubao-seed-1.6-250615)
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"API调用失败: {str(e)}")
            raise

    def process_chunks(self, chunks):
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            print(f"处理第 {i + 1}/{len(chunks)} 个文本块，字数: {len(chunk)}")
            try:
                processed_text = self.call_doubao_api(chunk)
                processed_chunks.append(processed_text)
                print(f"成功处理第 {i + 1} 个文本块")
            except Exception as e:
                print(f"处理第 {i + 1} 个文本块时出错: {str(e)}")
                processed_chunks.append(chunk)
            time.sleep(self.delay)
        return processed_chunks

    def create_txt(self, content):
        formatted_content = []
        for para in content:
            lines = para.split('\n')
            formatted_lines = ["　　" + line if line.strip() else line for line in lines]
            formatted_para = '\n'.join(formatted_lines)
            formatted_content.append(formatted_para)
        with open(self.txt_output_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(formatted_content))
        print(f"TXT文件已保存至: {self.txt_output_path}")

    def create_docx(self, content):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(12)
        style.paragraph_format.space_after = Pt(5)

        for para in content:
            lines = para.split('\n')
            doc_para = doc.add_paragraph()
            for line in lines:
                if line.strip():
                    doc_para.add_run(f"　　{line}\n")
                else:
                    doc_para.add_run("\n")
        doc.save(self.doc_output_path)
        print(f"Word文档已保存至: {self.doc_output_path}")

    def run(self):
        print("开始读取PDF文档...")
        paragraphs = self.read_pdf()
        print(f"成功读取 {len(paragraphs)} 个段落")

        print("开始创建文本块...")
        chunks = self.create_chunks(paragraphs)
        print(f"成功创建 {len(chunks)} 个文本块")

        print("开始调用API添加标点...")
        processed_chunks = self.process_chunks(chunks)
        print("所有文本块处理完成")

        print("开始生成TXT文件...")
        self.create_txt(processed_chunks)

        print("开始生成Word文档...")
        self.create_docx(processed_chunks)

        print("所有文件生成完成！")


if __name__ == "__main__":
    API_KEY = "api_key"                    ### 引号内替换为您的api密钥
    INPUT_PDF_PATH = r"___.pdf"            ### 引号内替换为导入的pdf文件地址
    TXT_OUTPUT_PATH = "___.txt"            ### 引号内替换为导出的txt文件名
    DOCX_OUTPUT_PATH = "___.docx"          ### 引号内替换为导出的doc文件名

    processor = PDFPunctuationProcessor(
        api_key=API_KEY,
        input_pdf_path=INPUT_PDF_PATH,
        txt_output_path=TXT_OUTPUT_PATH,
        doc_output_path=DOCX_OUTPUT_PATH,
        chunk_size=10000,
        #文本块大小，可结合模型能力自行修改，一般在1w-1.5w
        delay=2
        #api调用延迟，可自行修改，在不触发限流情况下以小为佳
    )
    processor.run()